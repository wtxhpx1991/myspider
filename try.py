# -*- coding:UTF-8 -*-
from urllib import request
from bs4 import BeautifulSoup
import re
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import qn

if __name__=="__main__":
    penalty_title_pattern = re.compile("[\u4e00-\u9fa5]+")
    # 设置头文件
    head = {}
    head[
        'User-Agent'] = "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36 SE 2.X MetaSr 1.0"

    # 初始网址
    start_download_url = "http://www.csrc.gov.cn/pub/zjhpublic/3300/3313/index_7401_"
    # 行政处罚决定书页面数量，目前是60页
    # 不知道怎么自动抓这个数字
    PAGE_NUMBER = 60
    # 存储位置
    FILEPATH = r"F:\book\行政处罚决定书"
    # 生成网页链接
    download_url_set = ["http://www.csrc.gov.cn/pub/zjhpublic/3300/3313/index_7401.htm"]
    for i in range(PAGE_NUMBER - 1):
        download_url_set.append(start_download_url + str(i + 1) + ".htm")

    # 对每一个页面，首先抓取全部链接，然后根据链接里的内容生成word
    print(datetime.now().strftime('%Y-%m-%d %H:%M:%S') + " 开始抓取")
    print("-----------------------------------------------------")
    for download_url in download_url_set:
        download_req = request.Request(url=download_url, headers=head)
        download_response = request.urlopen(download_req)
        download_html = download_response.read().decode('UTF-8', 'ignore')
        soup_texts = BeautifulSoup(download_html, 'lxml')
        # 获取页面核心数据
        soup_core_texts_html = soup_texts.find_all("div", attrs={"class": "row"})
        # 在每一页中提取处罚决定书年月以及网址后缀
        for tag in soup_core_texts_html:
            tag_number = tag.find("li", class_="wh").get_text()
            html_suffix = tag.find("a").attrs['href'].split("/")[-1]
            tag_datetime = tag.find("a").attrs['href'].split("/")[-2]
            sub_download_url = "http://www.csrc.gov.cn/pub/zjhpublic/G00306212/" + tag_datetime + "/" + html_suffix
            sub_download_req = request.Request(url=sub_download_url, headers=head)
            sub_download_response = request.urlopen(sub_download_req)
            sub_download_html = sub_download_response.read().decode('UTF-8', 'ignore')
            sub_soup_texts = BeautifulSoup(sub_download_html, 'lxml')
            penalty_title = sub_soup_texts.find_all("script")
            penalty_texts = sub_soup_texts.find_all("p", attrs={"class": "p0"})
            # 文档结果
            penalty_title_result = penalty_title_pattern.findall(str(penalty_title))
            # 新建word文档
            document = Document()
            title = document.add_heading(
                penalty_title_result[0] + "\n" + ",".join(penalty_title_result[1:]) + "\n" + penalty_texts[1].text,
                0)  # 插入标题
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
            title.style.font.color.rgb = RGBColor(255, 0, 0)
            for j in penalty_texts:
                p = document.add_paragraph(j.text)
                p.style.font.name = 'Times New Roman'
                p.style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.style.font.size = Pt(15)
            documenttitle = penalty_title_result[0] + tag_number + ",".join(penalty_title_result[1:]) + ".docx"
            document.save(FILEPATH + "\\" + documenttitle)
            print(datetime.now().strftime('%Y-%m-%d %H:%M:%S') + " 结束抓取" + documenttitle)
            print("-----------------------------------------------------")
    print(datetime.now().strftime('%Y-%m-%d %H:%M:%S') + " 结束抓取")





