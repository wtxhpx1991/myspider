# -*- coding:UTF-8 -*-
from urllib import request
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import qn
FILEPATH=r"F:\book\行政处罚决定书"

penalty_title_pattern = re.compile("[\u4e00-\u9fa5]+")

sub_download_url = "http://www.csrc.gov.cn/pub/zjhpublic/G00306212/201112/t20111227_204195.htm"
head = {}
head[
    'User-Agent'] = "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36 SE 2.X MetaSr 1.0"
sub_download_req = request.Request(url=sub_download_url, headers=head)
sub_download_response = request.urlopen(sub_download_req)
sub_download_html = sub_download_response.read().decode('UTF-8', 'ignore')
sub_soup_texts = BeautifulSoup(sub_download_html, 'lxml')
penalty_title = sub_soup_texts.find_all("script")
penalty_texts = sub_soup_texts.find_all("p", attrs={"style":"TEXT-INDENT: 25.2pt; LINE-HEIGHT: 27pt; MARGIN-RIGHT: 1pt"})
# penalty_texts = sub_soup_texts.find_all("div", class_="Custom_UnionStyle")

penalty_title_result = re.findall(r'"(.*?)"', str(penalty_title))[0]
penalty_texts_result = ""
for i in penalty_texts:
    penalty_texts_result = penalty_texts_result + i.text + "\n"
print(penalty_texts_result + penalty_texts_result)



document = Document()
title = document.add_heading(penalty_title_result[0] + "\n" + ",".join(penalty_title_result[1:])+ "\n" +penalty_texts[1].text, 0)  # 插入标题
title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
title.style.font.color.rgb=RGBColor(255,0,0)
for i in penalty_texts[4:len(penalty_texts)]:
    p = document.add_paragraph(i.text)
    p.style.font.name='Times New Roman'
    p.style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.style.font.size = Pt(15)
documenttitle= penalty_title_result[0] + penalty_texts[1].text+penalty_title_result[1]+".docx"
document.save(FILEPATH+"\\"+documenttitle)